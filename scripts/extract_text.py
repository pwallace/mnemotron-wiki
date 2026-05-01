# SPDX-License-Identifier: GPL-3.0-or-later
# Copyright (C) 2026 Patrick R. Wallace, Hamilton College LITS
# This file is part of Mnemotron Wiki. See LICENSE for terms.
#
# AI ASSISTANCE NOTICE: Developed with assistance from Claude (Anthropic).
# Reviewed and tested; verify behavior in your own environment.

"""
extract_text.py — Extract plain text from supported document types.

Supports: PDF, DOCX, plain text (.txt, .md), CSV, HTML/HTM.

Usage (command line):
    python scripts/extract_text.py path/to/document.pdf

Usage (from another script):
    from scripts.extract_text import extract_text
    result = extract_text(Path("path/to/document.pdf"))
    print(result["text"])

Return value is a dict:
    {
        "text":     str,        # extracted plain text
        "type":     str,        # detected file type (e.g. "pdf")
        "source":   Path,       # original file path
        "metadata": dict,       # any available metadata (title, author, etc.)
        "error":    str | None, # error message if extraction failed, else None
    }
"""

import csv
import sys
from io import StringIO
from pathlib import Path


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------

def _extract_pdf(filepath: Path) -> tuple[str, dict]:
    """Extract text from a PDF using pdfminer.six."""
    from pdfminer.high_level import extract_text as pdfminer_extract
    from pdfminer.pdfpage import PDFPage
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfparser import PDFParser

    text = pdfminer_extract(str(filepath))

    # Try to pull basic metadata from the PDF info dict.
    metadata = {}
    try:
        with open(filepath, "rb") as f:
            parser = PDFParser(f)
            doc = PDFDocument(parser)
            if doc.info:
                raw = doc.info[0]
                for key in ("Title", "Author", "Subject", "Keywords"):
                    value = raw.get(key)
                    if value:
                        # PDF metadata values are often bytes; decode safely.
                        if isinstance(value, bytes):
                            value = value.decode("utf-8", errors="replace")
                        metadata[key.lower()] = value.strip()
    except Exception:
        pass  # metadata is best-effort; don't fail extraction over it

    return text, metadata


def _extract_docx(filepath: Path) -> tuple[str, dict]:
    """Extract text from a .docx file using python-docx."""
    from docx import Document

    doc = Document(str(filepath))

    # Join all paragraph text; preserve blank lines between paragraphs.
    paragraphs = [para.text for para in doc.paragraphs]
    text = "\n\n".join(p for p in paragraphs if p.strip())

    # Core properties (author, title) if present.
    metadata = {}
    props = doc.core_properties
    if props.title:
        metadata["title"] = props.title
    if props.author:
        metadata["author"] = props.author
    if props.subject:
        metadata["subject"] = props.subject

    return text, metadata


def _extract_html(filepath: Path) -> tuple[str, dict]:
    """Extract visible text from an HTML file using BeautifulSoup."""
    from bs4 import BeautifulSoup

    raw = filepath.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "lxml")

    # Remove script and style elements — their text is not useful.
    for tag in soup(["script", "style", "head"]):
        tag.decompose()

    # get_text with a space separator, then tidy up blank lines.
    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(line for line in lines if line)

    # Try to extract <title> for metadata.
    metadata = {}
    title_tag = soup.find("title")
    if title_tag:
        metadata["title"] = title_tag.get_text(strip=True)

    return text, metadata


def _extract_csv(filepath: Path) -> tuple[str, dict]:
    """
    Represent a CSV as plain text.

    Each row becomes a line of pipe-separated values, making it readable
    as plain text while preserving structure.  The header row (if present)
    is preserved as-is.

    Pipe-separated format is chosen because it is human-readable, survives
    whitespace normalization, and is easy for Claude to parse into a table.
    """
    raw = filepath.read_text(encoding="utf-8", errors="replace")
    reader = csv.reader(StringIO(raw))
    rows = list(reader)

    lines = [" | ".join(cell.strip() for cell in row) for row in rows if any(row)]
    text = "\n".join(lines)

    metadata = {}
    if rows:
        metadata["columns"] = rows[0]  # first row treated as header
        metadata["row_count"] = str(len(rows) - 1)

    return text, metadata


def _extract_plaintext(filepath: Path) -> tuple[str, dict]:
    """Read a plain text or markdown file directly."""
    text = filepath.read_text(encoding="utf-8", errors="replace")
    return text, {}


def _extract_eml(filepath: Path) -> tuple[str, dict]:
    """Extract text and headers from an .eml email file."""
    import email
    from email import policy as email_policy

    msg = email.message_from_bytes(
        filepath.read_bytes(), policy=email_policy.default
    )

    metadata = {}
    for header in ("subject", "from", "to", "date"):
        value = msg.get(header)
        if value:
            metadata[header] = str(value).strip()

    # Collect text parts; prefer text/plain over text/html because plain
    # text is more compact and avoids re-parsing HTML artefacts.  Only
    # fall back to HTML if the message has no plain-text part at all
    # (common in heavily formatted marketing emails, rare in work mail).
    plain_parts: list[str] = []
    html_parts: list[str] = []
    for part in msg.walk():
        if str(part.get("Content-Disposition", "")).lower().startswith("attachment"):
            continue
        ct = part.get_content_type()
        if ct == "text/plain":
            plain_parts.append(part.get_content())
        elif ct == "text/html":
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(part.get_content(), "lxml")
            for tag in soup(["script", "style"]):
                tag.decompose()
            html_parts.append(soup.get_text(separator="\n"))

    parts = plain_parts if plain_parts else html_parts
    text = "\n\n".join(p.strip() for p in parts if p.strip())
    return text, metadata


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

# Map file extensions to extractor functions.
_EXTRACTORS = {
    ".pdf":  _extract_pdf,
    ".docx": _extract_docx,
    ".odt":  _extract_docx,   # python-docx handles ODT via the same OOXML API
    ".html": _extract_html,
    ".htm":  _extract_html,
    ".csv":  _extract_csv,
    ".txt":  _extract_plaintext,
    ".md":   _extract_plaintext,
    ".eml":  _extract_eml,
}


def extract_text(filepath: Path) -> dict:
    """
    Extract plain text from *filepath*.

    Returns a result dict as described in the module docstring.
    Never raises — errors are captured in result["error"].
    """
    filepath = Path(filepath)
    ext = filepath.suffix.lower()
    file_type = ext.lstrip(".")

    result = {
        "text": "",
        "type": file_type,
        "source": filepath,
        "metadata": {},
        "error": None,
    }

    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        result["error"] = f"Unsupported file type: {ext}"
        return result

    try:
        text, metadata = extractor(filepath)
        result["text"] = text.strip()
        result["metadata"] = metadata
    except Exception as exc:
        result["error"] = f"{type(exc).__name__}: {exc}"

    return result


# ---------------------------------------------------------------------------
# Command-line interface
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/extract_text.py <filepath>")
        sys.exit(1)

    target = Path(sys.argv[1])
    if not target.exists():
        print(f"Error: file not found: {target}")
        sys.exit(1)

    res = extract_text(target)

    if res["error"]:
        print(f"Extraction error: {res['error']}")
        sys.exit(1)

    if res["metadata"]:
        print("--- metadata ---")
        for k, v in res["metadata"].items():
            print(f"  {k}: {v}")
        print("--- text ---")

    print(res["text"])
